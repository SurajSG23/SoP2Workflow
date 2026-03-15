from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO

import pdfplumber
from docx import Document


@dataclass
class ParsedDocument:
    text: str


class DocumentParserError(Exception):
    """Raised when a document cannot be parsed."""


class DocumentParserService:
    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        suffix = filename.lower().rsplit(".", maxsplit=1)[-1]

        if suffix == "pdf":
            return self._parse_pdf(content)
        if suffix == "docx":
            return self._parse_docx(content)

        raise DocumentParserError("Unsupported file type. Please upload PDF or DOCX.")

    def _parse_pdf(self, content: bytes) -> ParsedDocument:
        chunks: list[str] = []
        with pdfplumber.open(BytesIO(content)) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text() or ""
                if extracted.strip():
                    chunks.append(extracted.strip())
        return ParsedDocument(text="\n".join(chunks))

    def _parse_docx(self, content: bytes) -> ParsedDocument:
        document = Document(BytesIO(content))
        chunks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        return ParsedDocument(text="\n".join(chunks))
